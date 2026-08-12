import os
import re
import base64
import json
import tempfile
import logging
from typing import Optional
from pathlib import Path

from fastapi import APIRouter, Form, UploadFile, File, HTTPException, Depends
from pydantic import BaseModel, Field
from docx import Document

from .auth import verify_api_key
from .database import (
    create_project,
    get_projects,
    get_project,
    update_project_tz,
    get_requirements_tree,
    delete_requirements_by_project,
    create_requirement,
)
from .main import client, MODEL, VISION_MODEL

router = APIRouter(prefix="/projects", tags=["projects"])

logger = logging.getLogger(__name__)
CHTZ_PROMPT_PATH = Path(__file__).resolve().parent / "prompts" / "chtz_generator.md"


def load_chtz_prompt() -> str:
    try:
        return CHTZ_PROMPT_PATH.read_text(encoding="utf-8")
    except FileNotFoundError:
        logger.warning("Файл скилла не найден: %s", CHTZ_PROMPT_PATH)
        return (
            "Ты — профессиональный системный аналитик. "
            "Извлеки из текста ТЗ функциональные требования в JSON-формате."
        )


CHTZ_PROMPT = load_chtz_prompt()


class ProjectCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=255)


class ProjectResponse(BaseModel):
    id: int
    name: str
    created_at: str


def extract_docx_text(document: Document) -> str:
    """Извлекает текст из параграфов и таблиц документа .docx."""
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_docx_images(document: Document) -> list[tuple[str, bytes]]:
    """Извлекает встроенные изображения из .docx и возвращает [(filename, bytes), ...]."""
    images = []
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                ext = image_part.content_type.split("/")[-1]
                if ext in ("png", "jpeg", "jpg", "gif", "bmp"):
                    images.append((f"image.{ext}", image_part.blob))
            except Exception as e:
                logger.warning("Не удалось извлечь изображение из .docx: %s", e)
    return images


async def describe_image(image_bytes: bytes) -> str:
    """Отправляет изображение в Vision-модель и возвращает описание."""
    image_data = base64.b64encode(image_bytes).decode("utf-8")
    mime = "image/png" if image_bytes[:8].startswith(b"\x89PNG") else "image/jpeg"
    response = await client.chat.completions.create(
        model=VISION_MODEL,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Кратко опиши изображение на русском языке. "
                            "Это иллюстрация из технического задания. "
                            "Перечисли ключевые элементы интерфейса или схемы."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{image_data}"},
                    },
                ],
            }
        ],
        max_tokens=1000,
        temperature=0.3,
    )
    return response.choices[0].message.content.strip()


def extract_chtz_json(raw_content: str) -> list[dict]:
    """Извлекает и валидирует JSON-структуру ЧТЗ из ответа модели."""
    try:
        parsed = json.loads(raw_content)
    except json.JSONDecodeError as e:
        raise ValueError(f"Ответ модели не является валидным JSON: {e}") from e

    if isinstance(parsed, list):
        sections = parsed
    elif isinstance(parsed, dict):
        sections = parsed.get("sections", [])
    else:
        raise ValueError("Ответ модели должен быть JSON-объектом или массивом")

    if not isinstance(sections, list):
        raise ValueError("Значение 'sections' должно быть массивом")

    if not sections:
        raise ValueError("Ответ модели не содержит секций")

    return sections


def _get_parent_code(code: str) -> Optional[str]:
    """Возвращает код родительского требования, отбросив последнюю часть после точки.

    Например, для 'ФТ_1.2.1' вернёт 'ФТ_1.2', для 'ФТ_1' вернёт None.
    """
    if not code or "." not in code:
        return None
    return code.rsplit(".", 1)[0]


def _extract_section_code(title: str, fallback: str) -> str:
    """Извлекает номер раздела из начала заголовка (например, '1.1.1').

    Если номер не найден, возвращает fallback.
    """
    match = re.match(r"(\d+(?:\.\d+)*)\.", title)
    if match:
        return match.group(1)
    return fallback


def _ensure_parent_requirement(
    project_id: int,
    parent_code: str,
    code_to_id: dict[str, int],
    section_id: int,
    section_path: str,
) -> int:
    """Возвращает id родительского требования, создавая заглушку при необходимости."""
    if parent_code in code_to_id:
        return code_to_id[parent_code]

    grandparent_code = _get_parent_code(parent_code)
    if grandparent_code:
        grandparent_id = _ensure_parent_requirement(
            project_id, grandparent_code, code_to_id, section_id, section_path
        )
    else:
        grandparent_id = section_id

    stub_id = create_requirement(
        project_id=project_id,
        title=parent_code,
        description="",
        code=parent_code,
        parent_id=grandparent_id,
        section_path=f"{section_path} > {parent_code}",
        sort_order=0,
    )
    code_to_id[parent_code] = stub_id
    return stub_id


def save_chtz_to_db(project_id: int, sections: list[dict]) -> None:
    """Сохраняет разделы и требования в таблицу requirements, выстраивая иерархию по кодам ФТ."""
    delete_requirements_by_project(project_id)

    for section_order, section in enumerate(sections):
        title = section.get("title", "Раздел")
        section_code = _extract_section_code(title, f"S_{section_order + 1}")
        section_path = title
        section_id = create_requirement(
            project_id=project_id,
            title=title,
            description="",
            code=section_code,
            parent_id=None,
            section_path=section_path,
            sort_order=section_order,
        )

        requirements = section.get("requirements", [])
        # Сортируем по глубине кода, чтобы родители обрабатывались раньше детей,
        # сохраняя исходный порядок внутри одного уровня.
        indexed_reqs = list(enumerate(requirements))
        indexed_reqs.sort(key=lambda item: item[1].get("code", "").count("."))

        code_to_id: dict[str, int] = {}
        for original_idx, req in indexed_reqs:
            code = req.get("code") or f"ФТ_{section_order + 1}.{original_idx + 1}"
            parent_code = _get_parent_code(code)

            if parent_code and parent_code in code_to_id:
                parent_id = code_to_id[parent_code]
            elif parent_code:
                parent_id = _ensure_parent_requirement(
                    project_id, parent_code, code_to_id, section_id, section_path
                )
            else:
                parent_id = section_id

            req_id = create_requirement(
                project_id=project_id,
                title=req.get("title", "Требование"),
                description=req.get("description", ""),
                code=code,
                parent_id=parent_id,
                section_path=f"{section_path} > {code}",
                sort_order=original_idx,
            )
            code_to_id[code] = req_id


@router.post("", response_model=ProjectResponse)
async def create_project_endpoint(
    project: ProjectCreate,
    api_key: str = Depends(verify_api_key),
):
    """Создать новый проект."""
    return create_project(project.name)


@router.get("", response_model=list[ProjectResponse])
async def list_projects(api_key: str = Depends(verify_api_key)):
    """Получить список проектов."""
    return get_projects()


@router.get("/{project_id}/tz")
async def get_project_tz(project_id: int, api_key: str = Depends(verify_api_key)):
    """Получить сохранённый текст ТЗ проекта."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")
    return {"project_id": project_id, "tz_text": project["tz_text"], "tz_filename": project["tz_filename"]}


@router.post("/{project_id}/upload-tz")
async def upload_tz(
    project_id: int,
    file: Optional[UploadFile] = File(None),
    tz_text: Optional[str] = Form(None),
    api_key: str = Depends(verify_api_key),
):
    """Загрузить техническое задание: файл .docx или текст."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")

    if not file and not tz_text:
        raise HTTPException(status_code=400, detail="Необходимо передать файл .docx или текст ТЗ")

    full_text = tz_text or ""
    tz_filename = None

    if file:
        tz_filename = file.filename
        suffix = Path(file.filename or "tz.docx").suffix.lower()
        if suffix not in (".docx", ".doc"):
            raise HTTPException(status_code=400, detail="Поддерживаются только файлы .docx")

        contents = await file.read()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(contents)
                tmp_path = tmp.name

            document = Document(tmp_path)
            doc_text = extract_docx_text(document)
            full_text = f"{full_text}\n\n{doc_text}".strip() if full_text else doc_text

            # Извлекаем и распознаём изображения из .docx
            images = extract_docx_images(document)
            if images:
                image_descriptions = []
                for idx, (filename, image_bytes) in enumerate(images, start=1):
                    try:
                        description = await describe_image(image_bytes)
                        image_descriptions.append(f"Изображение #{idx} ({filename}): {description}")
                    except Exception as e:
                        logger.warning("Ошибка распознавания изображения #%d: %s", idx, e)
                if image_descriptions:
                    full_text += "\n\n--- Описание изображений ---\n"
                    full_text += "\n\n".join(image_descriptions)
        except Exception as e:
            logger.error("Ошибка обработки .docx: %s", e)
            raise HTTPException(status_code=400, detail=f"Ошибка обработки файла .docx: {str(e)}")
        finally:
            if "tmp_path" in locals() and os.path.exists(tmp_path):
                os.remove(tmp_path)

    if not full_text.strip():
        raise HTTPException(status_code=400, detail="ТЗ пустое после обработки")

    updated = update_project_tz(project_id, full_text, tz_filename)
    return {
        "project_id": project_id,
        "tz_text": updated["tz_text"],
        "tz_filename": updated["tz_filename"],
    }


@router.post("/{project_id}/generate-chtz")
async def generate_chtz(
    project_id: int,
    api_key: str = Depends(verify_api_key),
):
    """Сгенерировать структуру ЧТЗ из ТЗ с помощью LLM и сохранить в БД."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")

    if not project["tz_text"].strip():
        raise HTTPException(status_code=400, detail="У проекта отсутствует загруженное ТЗ")

    try:
        response = await client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": CHTZ_PROMPT},
                {"role": "user", "content": f"Текст технического задания:\n\n{project['tz_text']}"},
            ],
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        raw_content = response.choices[0].message.content.strip()
        logger.info("Сырой ответ ЧТЗ: %s", raw_content[:1000])

        sections = extract_chtz_json(raw_content)
        save_chtz_to_db(project_id, sections)

        return {
            "project_id": project_id,
            "sections": sections,
            "requirements_count": sum(len(s.get("requirements", [])) for s in sections),
        }
    except HTTPException:
        raise
    except ValueError as e:
        logger.error("Ошибка парсинга ЧТЗ: %s", e)
        raise HTTPException(status_code=500, detail=f"Ошибка парсинга ЧТЗ: {e}")
    except Exception as e:
        logger.error("Ошибка генерации ЧТЗ: %s", e)
        raise HTTPException(status_code=500, detail=f"Ошибка генерации ЧТЗ: {str(e)}")


@router.get("/{project_id}/requirements-tree")
async def get_requirements_tree_endpoint(
    project_id: int,
    api_key: str = Depends(verify_api_key),
):
    """Получить дерево требований проекта."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")
    return {"project_id": project_id, "tree": get_requirements_tree(project_id)}
