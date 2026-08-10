import os
import re
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from .auth import verify_api_key
from .database import (
    get_all_pages,
    get_ft_requirements,
    get_pmi_testcases,
    get_project,
    get_requirements_coverage_matrix,
    get_template_by_type,
)

router = APIRouter(prefix="/projects", tags=["documents"])

# Директория для хранения собранных документов
DOCUMENTS_DIR = os.path.join(
    os.path.dirname(__file__), "static", "uploads", "documents"
)
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

# Плейсхолдеры, которые заменяются на таблицы
TABLE_PLACEHOLDERS = {"FT_TABLE", "TEST_CASES", "COVERAGE_MATRIX"}

# Плейсхолдеры, которые заменяются текстом
TEXT_PLACEHOLDERS = {"PAGES_DESCRIPTIONS", "REQUIREMENTS_LIST"}

# Все поддерживаемые плейсхолдеры
SUPPORTED_PLACEHOLDERS = TABLE_PLACEHOLDERS | TEXT_PLACEHOLDERS


class AssembleResponse(BaseModel):
    sections: list[dict]
    assembled_file_path: str


class SectionUpdate(BaseModel):
    content: str = Field(..., min_length=0)


class SectionResponse(BaseModel):
    type: str
    content: Any
    index: int


def _placeholder_pattern(name: str) -> str:
    """Возвращает регулярное выражение для плейсхолдера с учётом пробелов."""
    return r"\{\{\s*" + re.escape(name) + r"\s*\}\}"


def _has_placeholder(text: str, name: str) -> bool:
    return re.search(_placeholder_pattern(name), text) is not None


def _replace_placeholder(text: str, name: str, replacement: str) -> str:
    return re.sub(_placeholder_pattern(name), lambda _m: replacement, text)


def _build_ft_table(project_id: int) -> tuple[list[str], list[list[str]]]:
    """Формирует таблицу функциональных требований."""
    headers = ["Код ФТ", "Название", "Описание"]
    rows_data = get_ft_requirements(project_id)
    if not rows_data:
        return headers, [["—", "—", "—"]]
    rows = [
        [
            r.get("code") or "—",
            r.get("title", ""),
            r.get("description", ""),
        ]
        for r in rows_data
    ]
    return headers, rows


def _build_test_cases_table(project_id: int) -> tuple[list[str], list[list[str]]]:
    """Формирует таблицу тест-кейсов, включённых в ПМИ."""
    headers = ["Название", "Шаги", "Ожидаемый результат"]
    rows_data = get_pmi_testcases(project_id)
    if not rows_data:
        return headers, [["—", "—", "—"]]
    rows = [
        [tc["title"], tc["steps"], tc["expected_result"]] for tc in rows_data
    ]
    return headers, rows


def _build_coverage_matrix_table(project_id: int) -> tuple[list[str], list[list[str]]]:
    """Формирует матрицу покрытия требований."""
    headers = ["Код ФТ", "Кол-во чек-листов", "Кол-во тест-кейсов"]
    rows_data = get_requirements_coverage_matrix(project_id)
    if not rows_data:
        return headers, [["—", "0", "0"]]
    rows = [
        [
            r.get("code") or "—",
            str(r.get("checklists_count", 0)),
            str(r.get("testcases_count", 0)),
        ]
        for r in rows_data
    ]
    return headers, rows


def _build_table_for_placeholder(name: str, project_id: int) -> tuple[list[str], list[list[str]]]:
    """Возвращает заголовки и строки для табличного плейсхолдера."""
    if name == "FT_TABLE":
        return _build_ft_table(project_id)
    if name == "TEST_CASES":
        return _build_test_cases_table(project_id)
    if name == "COVERAGE_MATRIX":
        return _build_coverage_matrix_table(project_id)
    raise ValueError(f"Неизвестный табличный плейсхолдер: {name}")


def _build_pages_text(project_id: int) -> str:
    """Возвращает текстовое описание страниц."""
    pages = get_all_pages()
    if not pages:
        return ""
    return "\n".join(f"• {p['description']}" for p in pages)


def _build_requirements_text(project_id: int) -> str:
    """Возвращает нумерованный список требований."""
    requirements = get_ft_requirements(project_id)
    if not requirements:
        return ""
    lines = []
    for idx, req in enumerate(requirements, start=1):
        code = req.get("code") or "—"
        title = req.get("title", "")
        lines.append(f"{idx}. {code}: {title}")
    return "\n".join(lines)


def _build_text_for_placeholder(name: str, project_id: int) -> str:
    """Возвращает текстовое содержимое для плейсхолдера."""
    if name == "PAGES_DESCRIPTIONS":
        return _build_pages_text(project_id)
    if name == "REQUIREMENTS_LIST":
        return _build_requirements_text(project_id)
    raise ValueError(f"Неизвестный текстовый плейсхолдер: {name}")


def _find_placeholders(text: str) -> list[str]:
    """Находит все плейсхолдеры в тексте."""
    return re.findall(r"\{\{\s*([A-Z_]+)\s*\}\}", text)


def _insert_table_after(
    doc: Document,
    paragraph: Paragraph,
    headers: list[str],
    rows: list[list[str]],
) -> Table:
    """Вставляет таблицу после параграфа и удаляет сам параграф."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True

    # Заголовок
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)

    # Строки данных
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)

    # Перемещаем таблицу на место параграфа
    p_element = paragraph._element
    p_element.addnext(table._element)
    p_element.getparent().remove(p_element)
    return table


def _replace_paragraph_placeholders(doc: Document, paragraph: Paragraph, project_id: int) -> None:
    """Заменяет плейсхолдеры в параграфе на таблицы или текст."""
    text = paragraph.text
    placeholders = _find_placeholders(text)
    if not placeholders:
        return

    # Если параграф содержит только один табличный плейсхолдер,
    # заменяем его на полноценную таблицу docx.
    if len(placeholders) == 1:
        name = placeholders[0]
        if name in TABLE_PLACEHOLDERS:
            remaining = re.sub(_placeholder_pattern(name), "", text).strip()
            if not remaining:
                headers, rows = _build_table_for_placeholder(name, project_id)
                _insert_table_after(doc, paragraph, headers, rows)
                return

    # Обычная текстовая замена (табличные плейсхолдеры внутри текста удаляем)
    new_text = text
    for name in placeholders:
        if name in SUPPORTED_PLACEHOLDERS:
            if name in TABLE_PLACEHOLDERS:
                replacement = ""
            else:
                replacement = _build_text_for_placeholder(name, project_id)
        else:
            replacement = ""
        new_text = _replace_placeholder(new_text, name, replacement)
    paragraph.text = new_text


def _replace_cell_placeholders(doc: Document, project_id: int) -> None:
    """Заменяет плейсхолдеры внутри ячеек существующих таблиц."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    placeholders = _find_placeholders(text)
                    if not placeholders:
                        continue
                    new_text = text
                    for name in placeholders:
                        if name in SUPPORTED_PLACEHOLDERS:
                            # Внутри ячейки нельзя вставить отдельную таблицу,
                            # поэтому табличные плейсхолдеры заменяем пустой строкой.
                            replacement = (
                                _build_text_for_placeholder(name, project_id)
                                if name in TEXT_PLACEHOLDERS
                                else ""
                            )
                        else:
                            replacement = ""
                        new_text = _replace_placeholder(new_text, name, replacement)
                    paragraph.text = new_text


def _replace_document_placeholders(doc: Document, project_id: int) -> None:
    """Заменяет все плейсхолдеры в документе."""
    # Обрабатываем параграфы в обратном порядке, чтобы удаление/добавление
    # элементов не сбило индексацию.
    for paragraph in list(doc.paragraphs):
        _replace_paragraph_placeholders(doc, paragraph, project_id)

    # Заменяем плейсхолдеры в таблицах
    _replace_cell_placeholders(doc, project_id)


def _extract_sections(doc: Document) -> list[dict]:
    """Извлекает секции документа (заголовки, параграфы, таблицы) в порядке следования."""
    sections = []
    index = 0

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            style_name = ""
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name

            section_type = "heading" if style_name.startswith("Heading") else "paragraph"
            sections.append(
                {
                    "type": section_type,
                    "content": paragraph.text,
                    "index": index,
                }
            )
            index += 1

        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows[1:]
                ]
            else:
                headers = []
                rows = []
            sections.append(
                {
                    "type": "table",
                    "content": {"headers": headers, "rows": rows},
                    "index": index,
                }
            )
            index += 1

    return sections


def _assembled_file_path(project_id: int, doc_type: str) -> str:
    """Возвращает путь для сохранения собранного документа."""
    project_dir = os.path.join(DOCUMENTS_DIR, str(project_id))
    os.makedirs(project_dir, exist_ok=True)
    return os.path.join(project_dir, f"assembled_{doc_type}.docx")


@router.post("/{project_id}/documents/{doc_type}/assemble", response_model=AssembleResponse)
async def assemble_document(
    project_id: int,
    doc_type: str,
    api_key: str = Depends(verify_api_key),
):
    """Собрать документ по шаблону и данным проекта."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")

    template = get_template_by_type(project_id, doc_type)
    if not template:
        raise HTTPException(
            status_code=404,
            detail=f"Шаблон для типа '{doc_type}' не найден",
        )

    template_path = template["file_path"]
    if not os.path.exists(template_path):
        raise HTTPException(
            status_code=404,
            detail="Файл шаблона не найден на диске",
        )

    try:
        doc = Document(template_path)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Не удалось открыть шаблон .docx: {exc}",
        ) from exc

    _replace_document_placeholders(doc, project_id)

    assembled_path = _assembled_file_path(project_id, doc_type)
    try:
        doc.save(assembled_path)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Не удалось сохранить собранный документ: {exc}",
        ) from exc

    sections = _extract_sections(doc)
    return {"sections": sections, "assembled_file_path": assembled_path}


@router.get("/{project_id}/documents/{doc_type}/sections")
async def get_document_sections(
    project_id: int,
    doc_type: str,
    api_key: str = Depends(verify_api_key),
):
    """Получить секции ранее собранного документа."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")

    assembled_path = _assembled_file_path(project_id, doc_type)
    if not os.path.exists(assembled_path):
        raise HTTPException(
            status_code=404,
            detail="Сначала соберите документ",
        )

    try:
        doc = Document(assembled_path)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Не удалось открыть собранный документ: {exc}",
        ) from exc

    return {"sections": _extract_sections(doc)}


@router.put(
    "/{project_id}/documents/{doc_type}/sections/{section_index}",
    response_model=SectionResponse,
)
async def update_section(
    project_id: int,
    doc_type: str,
    section_index: int,
    body: SectionUpdate,
    api_key: str = Depends(verify_api_key),
):
    """Изменить текст параграфа или заголовка в собранном документе."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")

    assembled_path = _assembled_file_path(project_id, doc_type)
    if not os.path.exists(assembled_path):
        raise HTTPException(
            status_code=404,
            detail="Сначала соберите документ",
        )

    try:
        doc = Document(assembled_path)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Не удалось открыть собранный документ: {exc}",
        ) from exc

    current_index = 0
    target_element = None
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            if current_index == section_index:
                target_element = Paragraph(child, doc)
                break
            current_index += 1
        elif child.tag == qn("w:tbl"):
            if current_index == section_index:
                raise HTTPException(
                    status_code=400,
                    detail="Редактирование таблиц пока не поддерживается",
                )
            current_index += 1

    if target_element is None:
        raise HTTPException(status_code=404, detail="Секция не найдена")

    target_element.text = body.content

    style_name = ""
    if target_element.style and target_element.style.name:
        style_name = target_element.style.name
    section_type = "heading" if style_name.startswith("Heading") else "paragraph"

    try:
        doc.save(assembled_path)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Не удалось сохранить изменения: {exc}",
        ) from exc

    return SectionResponse(type=section_type, content=body.content, index=section_index)


@router.get("/{project_id}/documents/{doc_type}/export")
async def export_document(
    project_id: int,
    doc_type: str,
    api_key: str = Depends(verify_api_key),
):
    """Выгрузить собранный документ в формате .docx."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Проект не найден")

    assembled_path = _assembled_file_path(project_id, doc_type)
    if not os.path.exists(assembled_path):
        raise HTTPException(
            status_code=404,
            detail="Сначала соберите документ",
        )

    return FileResponse(
        path=assembled_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{doc_type}.docx",
    )
