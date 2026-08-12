import os
import json
import tempfile
import pytest
from unittest.mock import patch, MagicMock, AsyncMock
from fastapi.testclient import TestClient
from docx import Document


VALID_CHTZ_RESPONSE = {
    "sections": [
        {
            "title": "1. Авторизация",
            "requirements": [
                {
                    "code": "ФТ_1",
                    "title": "Реализовать авторизацию",
                    "description": "Пользователь должен иметь возможность авторизоваться в системе.",
                    "section": "1",
                },
                {
                    "code": "ФТ_1.1",
                    "title": "Вход по логину и паролю",
                    "description": "Пользователь может войти, указав логин и пароль.",
                    "section": "1",
                },
                {
                    "code": "ФТ_1.2",
                    "title": "Вход по SMS",
                    "description": "Пользователь может войти по одноразовому коду из SMS.",
                    "section": "1",
                },
            ],
        },
        {
            "title": "2. Переводы",
            "requirements": [
                {
                    "code": "ФТ_2",
                    "title": "Реализовать переводы",
                    "description": "Пользователь должен иметь возможность выполнять переводы.",
                    "section": "2",
                },
                {
                    "code": "ФТ_2.1",
                    "title": "Перевод между своими счетами",
                    "description": "Пользователь может перевести деньги между своими счетами.",
                    "section": "2",
                }
            ],
        },
    ]
}


@pytest.fixture(scope="module")
def client():
    import tempfile as tmp
    import backend.database as db_module

    tmp_db = tmp.mkstemp(suffix=".db")[1]
    db_module.DB_PATH = tmp_db
    db_module.init_db()

    from backend.main import app, limiter

    limiter.enabled = False

    client = TestClient(app)
    client.headers.update({"X-API-Key": "test-api-key-for-tests"})
    yield client

    if os.path.exists(tmp_db):
        os.unlink(tmp_db)


class TestCreateProject:
    def test_create_success(self, client):
        response = client.post("/projects", json={"name": "Тестовый проект"})
        assert response.status_code == 200
        data = response.json()
        assert data["name"] == "Тестовый проект"
        assert "id" in data
        assert "created_at" in data

    def test_create_empty_name_rejected(self, client):
        response = client.post("/projects", json={"name": ""})
        assert response.status_code == 422

    def test_create_missing_name_rejected(self, client):
        response = client.post("/projects", json={})
        assert response.status_code == 422


class TestListProjects:
    def test_list_projects(self, client):
        response = client.get("/projects")
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)
        assert len(data) >= 1
        assert "name" in data[0]


class TestGetProjectTz:
    def test_get_tz_existing_project(self, client):
        # Убедимся, что у первого проекта пустое ТЗ
        response = client.get("/projects")
        projects = response.json()
        project_id = projects[0]["id"]

        response = client.get(f"/projects/{project_id}/tz")
        assert response.status_code == 200
        data = response.json()
        assert data["project_id"] == project_id
        assert "tz_text" in data

    def test_get_tz_not_found(self, client):
        response = client.get("/projects/99999/tz")
        assert response.status_code == 404


class TestUploadTz:
    def test_upload_tz_text(self, client):
        response = client.post("/projects", json={"name": "Проект с текстовым ТЗ"})
        project_id = response.json()["id"]

        response = client.post(
            f"/projects/{project_id}/upload-tz",
            data={"tz_text": "Система должна позволять пользователю авторизоваться."},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["project_id"] == project_id
        assert "авторизоваться" in data["tz_text"]

    def test_upload_tz_text_and_docx(self, client):
        response = client.post("/projects", json={"name": "Проект с docx"})
        project_id = response.json()["id"]

        doc = Document()
        doc.add_heading("Техническое задание", level=1)
        doc.add_paragraph("Пользователь может создавать новые проекты.")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Поле"
        table.rows[0].cells[1].text = "Тип"

        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(tmp_fd)
        doc.save(tmp_path)

        with open(tmp_path, "rb") as f:
            response = client.post(
                f"/projects/{project_id}/upload-tz",
                files={"file": ("tz.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        os.unlink(tmp_path)

        assert response.status_code == 200
        data = response.json()
        assert data["project_id"] == project_id
        assert "создавать новые проекты" in data["tz_text"]
        assert "Поле" in data["tz_text"]
        assert data["tz_filename"] == "tz.docx"

    def test_upload_tz_no_file_and_no_text(self, client):
        response = client.post("/projects", json={"name": "Пустой проект"})
        project_id = response.json()["id"]

        response = client.post(f"/projects/{project_id}/upload-tz")
        assert response.status_code == 400
        assert "файл .docx или текст ТЗ" in response.json()["error"]

    def test_upload_tz_invalid_extension(self, client):
        response = client.post("/projects", json={"name": "Проект с pdf"})
        project_id = response.json()["id"]

        response = client.post(
            f"/projects/{project_id}/upload-tz",
            files={"file": ("tz.pdf", b"fake pdf", "application/pdf")},
        )
        assert response.status_code == 400
        assert ".docx" in response.json()["error"]

    def test_upload_tz_project_not_found(self, client):
        response = client.post(
            "/projects/99999/upload-tz",
            data={"tz_text": "Текст ТЗ"},
        )
        assert response.status_code == 404


class TestGenerateChtz:
    @patch("backend.projects.client.chat.completions.create", new_callable=AsyncMock)
    def test_generate_chtz_success(self, mock_create, client):
        response = client.post("/projects", json={"name": "Проект для ЧТЗ"})
        project_id = response.json()["id"]

        client.post(
            f"/projects/{project_id}/upload-tz",
            data={"tz_text": "1. Авторизация. Пользователь может войти по логину и паролю. Пользователь может войти по SMS. 2. Переводы. Пользователь может переводить деньги между счетами."},
        )

        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = json.dumps(VALID_CHTZ_RESPONSE, ensure_ascii=False)
        mock_create.return_value = mock_response

        response = client.post(f"/projects/{project_id}/generate-chtz")
        assert response.status_code == 200
        data = response.json()
        assert data["project_id"] == project_id
        assert data["requirements_count"] == 5

        # Проверим дерево
        tree_response = client.get(f"/projects/{project_id}/requirements-tree")
        assert tree_response.status_code == 200
        tree_data = tree_response.json()
        assert tree_data["project_id"] == project_id
        tree = tree_data["tree"]
        assert len(tree) == 2
        assert tree[0]["title"] == "1. Авторизация"
        assert tree[0]["code"] == "1"
        assert len(tree[0]["children"]) == 1
        assert tree[0]["children"][0]["code"] == "ФТ_1"
        assert len(tree[0]["children"][0]["children"]) == 2
        assert tree[0]["children"][0]["children"][0]["code"] == "ФТ_1.1"
        assert tree[1]["title"] == "2. Переводы"
        assert tree[1]["children"][0]["code"] == "ФТ_2"
        assert len(tree[1]["children"][0]["children"]) == 1
        assert tree[1]["children"][0]["children"][0]["code"] == "ФТ_2.1"

    @patch("backend.projects.client.chat.completions.create", new_callable=AsyncMock)
    def test_generate_chtz_array_response(self, mock_create, client):
        response = client.post("/projects", json={"name": "Проект с массивом ЧТЗ"})
        project_id = response.json()["id"]

        client.post(
            f"/projects/{project_id}/upload-tz",
            data={"tz_text": "1. Авторизация. Пользователь может войти по логину и паролю."},
        )

        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = json.dumps(
            VALID_CHTZ_RESPONSE["sections"], ensure_ascii=False
        )
        mock_create.return_value = mock_response

        response = client.post(f"/projects/{project_id}/generate-chtz")
        assert response.status_code == 200
        data = response.json()
        assert data["project_id"] == project_id
        assert data["requirements_count"] == 5
        assert len(data["sections"]) == 2

    @patch("backend.projects.client.chat.completions.create", new_callable=AsyncMock)
    def test_generate_chtz_missing_parent_stub(self, mock_create, client):
        response = client.post("/projects", json={"name": "Проект с пропущенным родителем"})
        project_id = response.json()["id"]

        client.post(
            f"/projects/{project_id}/upload-tz",
            data={"tz_text": "Некоторое техническое задание."},
        )

        chtz = {
            "sections": [
                {
                    "title": "1. Раздел",
                    "requirements": [
                        {
                            "code": "ФТ_1.1.1",
                            "title": "Дочернее требование",
                            "description": "Описание дочернего требования.",
                            "section": "1",
                        }
                    ],
                }
            ]
        }

        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = json.dumps(chtz, ensure_ascii=False)
        mock_create.return_value = mock_response

        response = client.post(f"/projects/{project_id}/generate-chtz")
        assert response.status_code == 200

        tree_response = client.get(f"/projects/{project_id}/requirements-tree")
        assert tree_response.status_code == 200
        tree = tree_response.json()["tree"]
        assert len(tree) == 1
        section = tree[0]
        assert len(section["children"]) == 1
        # Должна быть создана заглушка ФТ_1
        assert section["children"][0]["code"] == "ФТ_1"
        assert section["children"][0]["title"] == "ФТ_1"
        # Под ней заглушка ФТ_1.1
        assert len(section["children"][0]["children"]) == 1
        assert section["children"][0]["children"][0]["code"] == "ФТ_1.1"
        # Под ней исходное требование ФТ_1.1.1
        assert len(section["children"][0]["children"][0]["children"]) == 1
        assert section["children"][0]["children"][0]["children"][0]["code"] == "ФТ_1.1.1"

    @patch("backend.projects.client.chat.completions.create", new_callable=AsyncMock)
    def test_generate_chtz_empty_sections(self, mock_create, client):
        response = client.post("/projects", json={"name": "Проект с пустым ЧТЗ"})
        project_id = response.json()["id"]

        client.post(
            f"/projects/{project_id}/upload-tz",
            data={"tz_text": "Некоторое техническое задание."},
        )

        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = json.dumps({"sections": []}, ensure_ascii=False)
        mock_create.return_value = mock_response

        response = client.post(f"/projects/{project_id}/generate-chtz")
        assert response.status_code == 500
        assert "секций" in response.json()["error"]

    @patch("backend.projects.client.chat.completions.create", new_callable=AsyncMock)
    def test_generate_chtz_invalid_json(self, mock_create, client):
        response = client.post("/projects", json={"name": "Проект с плохим ЧТЗ"})
        project_id = response.json()["id"]

        client.post(
            f"/projects/{project_id}/upload-tz",
            data={"tz_text": "Некоторое техническое задание."},
        )

        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = "Это не JSON"
        mock_create.return_value = mock_response

        response = client.post(f"/projects/{project_id}/generate-chtz")
        assert response.status_code == 500
        assert "JSON" in response.json()["error"]

    def test_generate_chtz_no_tz(self, client):
        response = client.post("/projects", json={"name": "Проект без ТЗ"})
        project_id = response.json()["id"]

        response = client.post(f"/projects/{project_id}/generate-chtz")
        assert response.status_code == 400
        assert "ТЗ" in response.json()["error"]

    def test_generate_chtz_project_not_found(self, client):
        response = client.post("/projects/99999/generate-chtz")
        assert response.status_code == 404


class TestRequirementsTree:
    def test_requirements_tree_project_not_found(self, client):
        response = client.get("/projects/99999/requirements-tree")
        assert response.status_code == 404

    def test_requirements_tree_empty(self, client):
        response = client.post("/projects", json={"name": "Проект без требований"})
        project_id = response.json()["id"]

        response = client.get(f"/projects/{project_id}/requirements-tree")
        assert response.status_code == 200
        data = response.json()
        assert data["tree"] == []
