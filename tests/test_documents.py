import io
import os
import shutil
import pytest
from docx import Document
from fastapi.testclient import TestClient


@pytest.fixture(scope="module")
def client():
    """Создаёт тестовый клиент с временной БД."""
    import tempfile as tmp
    import backend.database as db_module

    tmp_db = tmp.mkstemp(suffix=".db")[1]
    db_module.DB_PATH = tmp_db
    db_module.init_db()

    from backend.main import app, limiter

    limiter.enabled = False

    test_client = TestClient(app)
    test_client.headers.update({"X-API-Key": "test-api-key-for-tests"})
    yield test_client

    if os.path.exists(tmp_db):
        os.unlink(tmp_db)


@pytest.fixture
def project(client):
    """Создаёт тестовый проект и чистит его файлы после теста."""
    response = client.post("/projects", json={"name": "DocFlow QA Project"})
    assert response.status_code == 200
    project_id = response.json()["id"]
    yield project_id

    # Удаляем связанные файлы и запись о проекте
    from backend.database import delete_project
    from backend.templates import TEMPLATE_UPLOAD_DIR
    from backend.documents import DOCUMENTS_DIR

    delete_project(project_id)

    for base_dir in (TEMPLATE_UPLOAD_DIR, DOCUMENTS_DIR):
        project_dir = os.path.join(base_dir, str(project_id))
        if os.path.exists(project_dir):
            shutil.rmtree(project_dir)


@pytest.fixture
def section_requirement(project):
    """Создаёт раздел и функциональное требование для проекта."""
    from backend.database import create_requirement

    section_id = create_requirement(
        project_id=project,
        title="1. Авторизация",
        parent_id=None,
        sort_order=0,
    )
    ft_id = create_requirement(
        project_id=project,
        title="Вход по логину и паролю",
        description="Пользователь может авторизоваться по логину и паролю.",
        code="ФТ_1.1",
        parent_id=section_id,
        sort_order=0,
    )
    return {
        "project_id": project,
        "section_id": section_id,
        "ft_id": ft_id,
    }


@pytest.fixture
def checklist_with_pmi_testcase(section_requirement):
    """Создаёт чек-лист и тест-кейс, включённый в ПМИ."""
    from backend.database import create_checklist_dict, create_testcase

    req = section_requirement
    checklist = create_checklist_dict(
        req["ft_id"],
        {
            "positive": [
                {"id": "p-1", "text": "Проверить успешную авторизацию"}
            ],
            "negative": [],
        },
    )
    testcase = create_testcase(
        title="Успешная авторизация по логину",
        steps="1. Открыть страницу входа\n2. Ввести логин и пароль\n3. Нажать 'Войти'",
        expected_result="Пользователь авторизован и перенаправлен на главную",
        checklist_id=checklist["id"],
        requirement_id=req["ft_id"],
        include_in_pmi=True,
    )
    return {
        "project_id": req["project_id"],
        "ft_id": req["ft_id"],
        "checklist_id": checklist["id"],
        "testcase_id": testcase["id"],
    }


@pytest.fixture
def template_no_placeholders(tmp_path):
    """Возвращает путь к .docx шаблону без плейсхолдеров."""
    doc = Document()
    doc.add_heading("Документ без плейсхолдеров", level=1)
    doc.add_paragraph("Просто текст.")
    path = tmp_path / "template_no_ph.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def template_with_placeholders(tmp_path):
    """Возвращает путь к .docx шаблону с плейсхолдерами."""
    doc = Document()
    doc.add_heading("Программа и методика испытаний", level=1)
    doc.add_paragraph("Функциональные требования:")
    doc.add_paragraph("{{FT_TABLE}}")
    doc.add_paragraph("Тест-кейсы:")
    doc.add_paragraph("{{TEST_CASES}}")
    path = tmp_path / "template_ph.docx"
    doc.save(str(path))
    return str(path)


def _upload_template(client, project_id, file_path, doc_type):
    """Вспомогательная функция для загрузки шаблона."""
    with open(file_path, "rb") as f:
        response = client.post(
            f"/projects/{project_id}/templates",
            data={"doc_type": doc_type},
            files={
                "file": (
                    "template.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    return response


class TestTemplateUpload:
    def test_upload_template_success(self, client, project, template_no_placeholders):
        response = _upload_template(client, project, template_no_placeholders, "ПМИ")
        assert response.status_code == 200
        data = response.json()
        assert data["doc_type"] == "ПМИ"
        assert "id" in data
        assert "file_path" in data
        assert os.path.exists(data["file_path"])

    def test_upload_template_duplicate_doc_type(self, client, project, template_no_placeholders):
        _upload_template(client, project, template_no_placeholders, "ПМИ")
        response = _upload_template(client, project, template_no_placeholders, "ПМИ")
        assert response.status_code == 409
        assert "уже существует" in response.json()["error"]


class TestTemplateList:
    def test_list_templates(self, client, project, template_no_placeholders):
        _upload_template(client, project, template_no_placeholders, "ПЗ")
        response = client.get(f"/projects/{project}/templates")
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)
        assert len(data) == 1
        assert data[0]["doc_type"] == "ПЗ"
        assert "file_path" in data[0]


class TestTemplateDelete:
    def test_delete_template(self, client, project, template_no_placeholders):
        upload_resp = _upload_template(client, project, template_no_placeholders, "ПМИ")
        file_path = upload_resp.json()["file_path"]

        response = client.delete(f"/projects/{project}/templates/ПМИ")
        assert response.status_code == 200
        assert response.json()["doc_type"] == "ПМИ"
        assert not os.path.exists(file_path)

        response = client.delete(f"/projects/{project}/templates/ПМИ")
        assert response.status_code == 404


class TestDocumentAssemble:
    def test_assemble_no_placeholders(self, client, project, template_no_placeholders):
        _upload_template(client, project, template_no_placeholders, "Док")
        response = client.post(f"/projects/{project}/documents/Док/assemble")
        assert response.status_code == 200
        data = response.json()
        assert "sections" in data
        assert "assembled_file_path" in data
        assert os.path.exists(data["assembled_file_path"])

        texts = [s["content"] for s in data["sections"] if s["type"] == "paragraph"]
        assert "Просто текст." in texts

        # Проверяем, что в собранном файле нет незаменённых плейсхолдеров
        doc = Document(data["assembled_file_path"])
        for paragraph in doc.paragraphs:
            assert "{{" not in paragraph.text

    def test_assemble_with_placeholders(
        self, client, checklist_with_pmi_testcase, template_with_placeholders
    ):
        project_id = checklist_with_pmi_testcase["project_id"]
        _upload_template(client, project_id, template_with_placeholders, "ПМИ")
        response = client.post(f"/projects/{project_id}/documents/ПМИ/assemble")
        assert response.status_code == 200
        data = response.json()
        sections = data["sections"]
        assert "assembled_file_path" in data

        # Проверяем, что в JSON-секциях появились подставленные таблицы
        table_rows = []
        for section in sections:
            if section["type"] == "table":
                table_rows.extend(section["content"]["rows"])

        assert any("ФТ_1.1" in cell for row in table_rows for cell in row)
        assert any("Успешная авторизация по логину" in cell for row in table_rows for cell in row)

        # Проверяем, что в .docx файле плейсхолдеры заменены
        doc = Document(data["assembled_file_path"])
        for paragraph in doc.paragraphs:
            assert "{{FT_TABLE}}" not in paragraph.text
            assert "{{TEST_CASES}}" not in paragraph.text

        # Проверяем, что в таблицах файла есть данные
        file_rows = []
        for table in doc.tables:
            for row in table.rows:
                file_rows.append([cell.text.strip() for cell in row.cells])
        assert any("ФТ_1.1" in cell for row in file_rows for cell in row)
        assert any("Успешная авторизация по логину" in cell for row in file_rows for cell in row)

    def test_assemble_nonexistent_doc_type(self, client, project):
        response = client.post(f"/projects/{project}/documents/НетТакого/assemble")
        assert response.status_code == 404


class TestDocumentSections:
    def test_get_sections_without_assemble(self, client, project):
        response = client.get(f"/projects/{project}/documents/ПМИ/sections")
        assert response.status_code == 404
        assert "Сначала соберите" in response.json()["error"]

    def test_get_sections_after_assemble(self, client, project, template_no_placeholders):
        _upload_template(client, project, template_no_placeholders, "Док")
        assemble_resp = client.post(f"/projects/{project}/documents/Док/assemble")
        assert assemble_resp.status_code == 200

        response = client.get(f"/projects/{project}/documents/Док/sections")
        assert response.status_code == 200
        assert "sections" in response.json()


class TestUpdateSection:
    def test_update_section_and_export(self, client, project, template_no_placeholders):
        _upload_template(client, project, template_no_placeholders, "Док")
        assemble_resp = client.post(f"/projects/{project}/documents/Док/assemble")
        assert assemble_resp.status_code == 200
        sections = assemble_resp.json()["sections"]
        assembled_path = assemble_resp.json()["assembled_file_path"]

        # Находим индекс параграфа с исходным текстом
        section_index = None
        for section in sections:
            if section["type"] == "paragraph" and section["content"] == "Просто текст.":
                section_index = section["index"]
                break
        assert section_index is not None

        new_text = "Обновлённый текст секции."
        response = client.put(
            f"/projects/{project}/documents/Док/sections/{section_index}",
            json={"content": new_text},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["content"] == new_text
        assert data["type"] == "paragraph"
        assert data["index"] == section_index

        # Проверяем, что в файле появился новый текст
        doc = Document(assembled_path)
        paragraph_texts = [p.text for p in doc.paragraphs]
        assert new_text in paragraph_texts

        # Проверяем экспорт
        export_resp = client.get(f"/projects/{project}/documents/Док/export")
        assert export_resp.status_code == 200
        assert "application/vnd.openxmlformats" in export_resp.headers["content-type"]

        import io
        exported_doc = Document(io.BytesIO(export_resp.content))
        assert any(new_text in p.text for p in exported_doc.paragraphs)


class TestDocumentExport:
    def test_export_document(self, client, project, template_no_placeholders):
        _upload_template(client, project, template_no_placeholders, "Док")
        client.post(f"/projects/{project}/documents/Док/assemble")

        response = client.get(f"/projects/{project}/documents/Док/export")
        assert response.status_code == 200
        assert "application/vnd.openxmlformats" in response.headers["content-type"]
        assert response.content.startswith(b"PK")

    def test_export_without_assemble(self, client, project):
        response = client.get(f"/projects/{project}/documents/Док/export")
        assert response.status_code == 404
